import csv
from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = BASE_DIR / "generated_dataset" / "docx"
INDEX_FILE = BASE_DIR / "generated_dataset" / "dataset_index.csv"


DOCUMENT_TYPES = [
    "Drawing",
    "Specification",
    "Method Statement",
    "Material Submittal",
    "Shop Drawing",
    "Inspection Report",
    "Contract",
    "Meeting Minutes",
    "RFI"
]


DISCIPLINES = [
    "Architectural",
    "Structural",
    "Mechanical",
    "Electrical",
    "Civil"
]


KEYWORDS = {
    "Drawing": "drawing plan section elevation layout structural drawing architectural drawing",
    "Specification": "specification technical specification material standard performance requirement",
    "Method Statement": "method statement work procedure construction methodology sequence of work",
    "Material Submittal": "material submittal manufacturer datasheet catalogue material approval",
    "Shop Drawing": "shop drawing fabrication installation detail coordination drawing",
    "Inspection Report": "inspection report inspection result site inspection approved with comments",
    "Contract": "contract agreement conditions of contract scope of work payment terms",
    "Meeting Minutes": "meeting minutes minutes of meeting attendees agenda action items",
    "RFI": "rfi request for information clarification response required"
}


def create_sample_document(index: int, document_type: str):

    discipline = DISCIPLINES[index % len(DISCIPLINES)]

    title = f"{discipline} {document_type} Package {index:03d}"

    filename = f"DOC-{index:03d}_{document_type.replace(' ', '_')}.docx"

    document = Document()

    document.add_heading(title, level=1)

    document.add_paragraph(f"Document Title: {title}")
    document.add_paragraph(f"Revision Number: Rev. {chr(65 + index % 5)}")
    document.add_paragraph("Project Name: Cubic Tower Development")
    document.add_paragraph("Contractor: BuildRight Contracting LLC")
    document.add_paragraph("Consultant: Cubic Engineering Consultancy")
    document.add_paragraph(f"Submission Date: 2026-06-{(index % 28) + 1:02d}")
    document.add_paragraph(f"Discipline: {discipline}")
    document.add_paragraph("")
    document.add_paragraph(f"Document Type Keywords: {KEYWORDS[document_type]}")
    document.add_paragraph(
        "This file is a generated sample document for testing classification "
        "and metadata extraction in the AI Document Control Assistant."
    )

    output_path = OUTPUT_DIR / filename

    document.save(output_path)

    return {
        "filename": filename,
        "expected_type": document_type,
        "title": title,
        "discipline": discipline
    }


def main():

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    rows = []

    for index in range(1, 51):
        document_type = DOCUMENT_TYPES[(index - 1) % len(DOCUMENT_TYPES)]
        rows.append(
            create_sample_document(index, document_type)
        )

    with open(INDEX_FILE, "w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(
            file,
            fieldnames=[
                "filename",
                "expected_type",
                "title",
                "discipline"
            ]
        )

        writer.writeheader()
        writer.writerows(rows)

    print(f"Generated {len(rows)} sample documents.")
    print(f"Dataset saved to: {OUTPUT_DIR}")
    print(f"Index saved to: {INDEX_FILE}")


if __name__ == "__main__":
    main()