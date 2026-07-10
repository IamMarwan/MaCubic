from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parents[1]
SAMPLE_DIR = BASE_DIR / "data" / "samples"


def create_docx(filename, lines):
    doc = Document()

    doc.add_heading(lines[0], level=1)

    for line in lines[1:]:
        if line.endswith(":"):
            doc.add_heading(line[:-1], level=2)
        else:
            doc.add_paragraph(line)

    doc.save(SAMPLE_DIR / filename)


def main():
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)

    create_docx(
        "complete_drawing.docx",
        [
            "Architectural Drawing",
            "Project Name: Cubic Tower Development",
            "Document No: DWG-ARC-001",
            "Revision: A",
            "Date: 2026-07-10",
            "Approval Status: Approved",
            "Discipline: Architectural",
            "Prepared By: Design Team",
            "Checked By: Senior Architect",
            "Approved By: Project Manager",
            "Title Block:",
            "Drawing Number: DWG-ARC-001",
            "Scale: 1:100",
            "Revision History:",
            "Revision A issued for approval.",
            "Approval:",
            "Approved for construction.",
            "Notes:",
            "All dimensions must be verified on site.",
            "Signature: Prepared By",
            "Signature: Checked By",
            "Signature: Approved By",
        ],
    )

    create_docx(
        "incomplete_method_statement.docx",
        [
            "Method Statement",
            "Project Name: Cubic Tower Development",
            "Document No: MS-CIV-002",
            "Date: 2026-07-10",
            "Prepared By: Site Engineer",
            "Scope:",
            "Concrete casting works.",
            "Procedure:",
            "Follow approved drawings.",
            "Safety:",
            "Risk assessment before work.",
        ],
    )

    create_docx(
        "material_submittal_missing_signature.docx",
        [
            "Material Submittal",
            "Project Name: Cubic Tower Development",
            "Document No: MAT-MEP-003",
            "Revision: B",
            "Date: 2026-07-10",
            "Approval Status: Under Review",
            "Material Description:",
            "HVAC insulation material.",
            "Manufacturer:",
            "ABC Manufacturing",
            "Technical Data:",
            "Datasheet attached.",
            "Compliance Statement:",
            "Complies with project specification.",
            "Approval:",
            "Submitted for review.",
            "Revision History:",
            "Revision B updated.",
            "Signature: Prepared By",
        ],
    )

    print(f"Sample documents created in: {SAMPLE_DIR}")


if __name__ == "__main__":
    main()